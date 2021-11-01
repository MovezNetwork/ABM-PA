import networkx as nx
import json
import numpy as np
import os
import pandas as pd
import time
from docx import Document 
from xlwt import Workbook
import collections
from docx.shared import Inches
import matplotlib.pyplot as plt
from scipy.stats import pearsonr
import seaborn as sns
from itertools import chain

import newsrc.population as p
import newsrc.model as m

'''
TODO: should we include a "reset"(= clean up some objects; so you can restart a run without running init again) & "stop" (= clean up all objects) method

REPLY: as discussed, we should read about Python's garbage collector and decide on this
'''
    
class Simulation:
    def __init__(self, **args):
        self.input_args = self.load_input_args()
        self.PeerNominatedDataPopulation = p.PeerNominatedDataPopulation('Peer-Nominated data population', self.input_args)
        self.CommunicationDataPopulation = p.CommunicationDataPopulation('Communication data population', self.input_args)
        self.model = m.DiffusionModel('Gabrianelli Diffusion Model', self.input_args)


    def load_input_args(self):
        try:
            input_args = json.loads(open('../input/simulation.json').read())
        except Exception as ex:
            print('simulation.json does not exist!')
            print(ex)
            
        return input_args
    
    
    def simulate_interventions(self,time,population_name):
        
        percent = self.input_args['percent'] 
        generateGephiFiles = self.input_args['generateGephiFiles'] 
        writeToExcel = self.input_args['writeToExcel'] 
        intervention_strategies = self.input_args['intervention_strategy']

        #selected infuential agents
        simulation_selected_agents = {}
        #outcomes of the intervention
        simulation_outcomes_child = {}
        simulation_outcomes_avg = {}
        
        if(population_name == 'peer'):
            population = self.PeerNominatedDataPopulation
        elif(population_name == 'communication'):
            population = self.CommunicationDataPopulation
        
        for classroom_population in population.get_class_graphs(population.graph):

            classroom_population_id = list(classroom_population.nodes(data='class'))[1][1]
            simulation_outcomes_child[str(classroom_population_id)] = {}
            simulation_outcomes_avg[str(classroom_population_id)] = {}
            simulation_selected_agents[str(classroom_population_id)] = {}
            
            

            for intervention in intervention_strategies:
                c = classroom_population.copy()
                # modifies the PA of particular agents selected as influential 
                agent_selection_tuple = population.select_influential_agents(c, percent, intervention, False)
                # updated graph with enhanced PA in influential agents
                cl_pop = agent_selection_tuple[0]
                selected_agents = agent_selection_tuple[1]
                #running the simulation, executing the model with every timestamp
                for t in range(0,time):
                    cl_pop = self.model.execute(cl_pop,t)
                
                outcomes_in_dict = self.get_intervention_PA_dictionary(cl_pop)
                simulation_outcomes_child[str(classroom_population_id)][intervention] = outcomes_in_dict
                simulation_outcomes_avg[str(classroom_population_id)][intervention] = outcomes_in_dict.mean(axis=1)
                simulation_selected_agents[str(classroom_population_id)][intervention] = selected_agents
        
        if writeToExcel:
            self.interventionResultsToExcel(simulation_outcomes_child)
            
        df_agents_list = []
        for outer_dict in simulation_selected_agents.items():
            for intv in outer_dict[1]:
                if(intv != 'nointervention'):
                    df_agents_list.append([outer_dict[0],intv,outer_dict[1][intv]])

        df_agents = pd.DataFrame(df_agents_list, columns = ["SchoolClass", "Intervention", "InfluenceAgents"])
        df_agents.to_excel('selected_agents_'+ population_name +'.xlsx')  


        
        return simulation_outcomes_child,simulation_outcomes_avg,df_agents
            
    def get_intervention_PA_dictionary(self,graph):
        results_dict = dict(graph.nodes(data=True))
        PA_dict = {}
        for k, v in results_dict.items():
            PA_dict[k] = results_dict[k]['PA_hist']
        
                
        return pd.DataFrame(PA_dict)
    
        
    def interventionResultsToExcel(self,results):
        # loop the classes
        for class_id,res in results.items():
            
            directory='../output/Class'+repr(int(float(class_id)))
            
            if not os.path.exists(directory):
                os.makedirs(directory)
            
            filename=directory+'/'+repr(int(float(class_id)))+'.xls'
            
            w = Workbook()
            ws = w.add_sheet(class_id.strip('\'') + ' Simulation Outcomes')
                          
            #loop the dataframe
            rowNum=0
            colNum=0  
            
            for i in self.input_args['intervention_strategy']:                             
                col=list(res[i])
                colLen=len(col)
                
                #writing the labels in excel sheet
                ws.write(rowNum,colNum,i)
                for c in col: 
                    rowNum=rowNum+1
                    ws.write(rowNum,colNum,c)

                colNum = colNum + 1
                rowNum = 0

            w.save(filename)
            
    def getSuccessRates(self,results):

        success_rates = []
        for class_id,res in results.items():
            for i in self.input_args['intervention_strategy']:
                outcome = res[i]
                success_rates.append([class_id, i, self.get_change(outcome[364],outcome[0]), outcome[0] , outcome[364]])
            
        success_rates = pd.DataFrame(success_rates, columns = ["SchoolClass", "Intervention", "SuccessRate", "StartIntervention", "EndIntervention"])
        
        
        if(self.input_args['writeToExcel'] == True):
            directory='../output/ClassesSummary'
            if not os.path.exists(directory):
                os.makedirs(directory)
            filename=directory+'/SuccessRates.xls'
            writer = pd.ExcelWriter(filename)
            success_rates.to_excel(writer,'InterventionDif')
            writer.save()
        
        return success_rates

    def get_change(self,current, previous):
        if current == previous:
            return 0
        try:
            return round(((current - previous)/previous)*100.0,2)
        except ZeroDivisionError:
            return 0 
        
        
        
    
    def plot_interventions_per_participant(self,results):    

        '''

        Saves (Displays) plots of intervention results per children per class, based on the dictionary input. 
            results - dictionary containing the interventions' results applied per class
        '''

        for class_id,res in results.items():            

            directory='../output/Class'+class_id
            if not os.path.exists(directory):
                os.makedirs(directory)

            document = Document()
            filename=directory+'/'+'Class'+class_id+'_Interventions_Detailed_Per_Child.docx'
            document.add_paragraph('All Intervention plots for class ' + class_id + ' per participant')
                
            for i in self.input_args['intervention_strategy']:
                ax = res[i].plot(figsize=((15,10)))
                ax.legend(loc="upper right")
                ax.set_title('Class:'+ class_id+ ' Intervention:'+ i)
                fig =ax.get_figure()
                titlefig=directory +'/'+class_id.strip('\'')+i.strip('\'')+'.png'
                fig.savefig(titlefig)
                document.add_picture(titlefig, width=Inches(7))
                if not self.input_args['save_png']:
                    os.remove(titlefig)

            document.save(filename)
            
            
    def plot_interventions_per_class(self,results_avg):
        for class_id,res in results_avg.items():
            plt.figure(figsize=((15,10)))
            plt.xlim(0,364)
            plt.xlabel('Days')
            plt.ylabel('Mean PA')
            j = 0
            for i in self.input_args['intervention_strategy']:
                results_avg[class_id][i].plot(color=self.input_args['intervention_color'][j],label= i)
                j = j + 1
            plt.legend(title='All Interventions '+  class_id, loc="upper right")
            
            
    def plot_interventions_averaged(self,results_avg):
        all_averaged = {}
        for i in self.input_args['intervention_strategy']:
            temp_res = pd.Series([], dtype = float)
            counter = 0
            for class_id,res in results_avg.items():
                temp_res = temp_res.add(results_avg[class_id][i],fill_value=0)
                counter = counter + 1
            all_averaged[i] = temp_res/counter

        plt.figure(figsize=((15,10)))
        plt.xlim(0,364)
        plt.xlabel('Days')
        plt.ylabel('Mean PA')
        j = 0
        for i in self.input_args['intervention_strategy']:
            all_averaged[i].plot(color=self.input_args['intervention_color'][j],label= i)
            j = j + 1
        plt.legend(title='All Interventions ', loc="upper right")
            
            
    def heatmap(self,success_rates):
        
        success_rates  = success_rates.groupby(['SchoolClass'])['SuccessRate'].mean().reset_index()
        hm = success_rates[['SchoolClass','SuccessRate']].sort_values('SuccessRate',ascending=False)
        hm=hm.drop(hm.index[len(hm)-1])
        yrows=[1,1,1,1,1,2,2,2,2,2,3,3,3,3,3,4,4,4,4,4,5,5,5,5,5]
        hm=hm.assign(Yrows=yrows)
        xrows=[1,2,3,4,5,1,2,3,4,5,1,2,3,4,5,1,2,3,4,5,1,2,3,4,5]
        hm=hm.assign(Xrows=xrows)
        
        writer = pd.ExcelWriter('../output/heatmap.xlsx')
        hm.to_excel(writer, 'heatmap')
        
        cls=(np.asarray(hm['SchoolClass'])).reshape(5,5)
        perc=(np.asarray(hm['SuccessRate'])).reshape(5,5)
        result=hm.pivot(index='Yrows',columns='Xrows',values='SuccessRate')
        labels=(np.asarray(["Class {0} \n \n {1:.2f}%".format(c,p)
                          for c,p in zip(cls.flatten(),perc.flatten())])
               ).reshape(5,5)
        fig, ax =plt.subplots(figsize=(15,10))
        title= "Success Rates per Class"
        plt.title(title,fontsize=18)
        ttl=ax.title
        ttl.set_position([0.5,1.05])
        ax.set_xticks([])
        ax.set_yticks([])
        ax.axis('off')
        sns.heatmap(result,annot=labels,annot_kws={"size": 16},fmt="",cmap="Blues",linewidths=0.30,ax=ax)
        plt.show()