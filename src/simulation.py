
import networkx as nx
import numpy as np
import os
import pandas as pd
from docx import Document
from xlwt import Workbook
from docx.shared import Inches
import matplotlib.pyplot as plt
import seaborn as sns

import src.population as p
import src.model as m
import src.utils as utils

    
class Simulation:
    '''
        Simulation Class.
    '''    
    def __init__(self, **args):
        self.input_args = utils.load_input_args(file = '../input/simulation.json')
        self.PeerNominatedDataPopulation = p.PeerNominatedDataPopulation('Peer-Nominated data population', self.input_args)
        self.CommunicationDataPopulation = p.CommunicationDataPopulation('Communication data population', self.input_args)
        self.model = m.DiffusionModel(self.input_args)
        
    def simulate_preinterventions(self,time,population_name,threshold,ipa):
        '''
        Method for running the ABM pre-simulation (model calibration).
        Args:
            time (Integer): timepoint of a simulation (in days)
            population_name (str): string constant for selecting a particular population
            threshold (float): diffusion model thres_PA
            ipa(float): diffusion model I_PA
        Returns:
            dictionary: Simulation outcomes dictionaries. Full (per child) and averaged (per class). Followed by a dataframe with the selected influencers per class
        '''

        # set model parameters
        self.model.setThresholdPA(threshold)
        self.model.setIPA(ipa)

        if(population_name == 'nomination'):
            population = self.PeerNominatedDataPopulation
        elif(population_name == 'communication'):
            population = self.CommunicationDataPopulation

        # outcomes of the intervention
        simulation_outcomes_child = {}
        simulation_outcomes_class = {}
        snapshot_pa = {}

        # run the model per class. Each class_population is a graph
        for class_population in population.get_class_graphs(population.graph):
            class_id = list(class_population.nodes(data='class'))[1][1]
            simulation_outcomes_child[str(class_id)] = {}
            simulation_outcomes_class[str(class_id)] = {}

            cl_pop = class_population.copy()
            #print('before preintervention class', c.nodes.data())
            #running the simulation, executing the model with every timestamp
            for t in range(0,time):
                sim_pop = self.model.execute(cl_pop,t)

            outcomes_in_dict = utils.get_PA_dictionary(sim_pop)

            #snapshot of the last 
            results_dict = dict(sim_pop.nodes(data=True))
            for k, v in results_dict.items():
                #taking the last element of all the saved PA_hist values
                snapshot_pa[k] = results_dict[k]['PA_hist'][-1]
                #print('after preintervention',snapshot_pa[k])
               
            simulation_outcomes_child[str(class_id)] = outcomes_in_dict
            simulation_outcomes_class[str(class_id)] = outcomes_in_dict.mean(axis=1)

        return simulation_outcomes_child, simulation_outcomes_class, snapshot_pa


    def simulate_interventions(self,time,population_name,preintervention_snapshot_PA,threshold,ipa,run_no):
        '''
        Method for running the ABM simulations.

        Args:
            time (Integer): timepoint of a simulation (in days)
            population_name (str): string constant for selecting a particular population
            threshold (float): diffusion model thres_PA
            ipa(float): diffusion model I_PA

        Returns:
            dictionary: Simulation outcomes dictionaries. Full (per child) and averaged (per class). Followed by a dataframe with the selected influencers per class
        '''        

        generateGephiFiles = self.input_args['generateGephiFiles'] 
        writeToExcel = self.input_args['writeToExcel']

        # intervention parameters
        intervention_strategies = self.input_args['intervention_strategy']
        percent = self.input_args['percent']

        # set model paramters
        self.model.setThresholdPA(threshold)
        self.model.setIPA(ipa)

        # outcome: selected influential agents
        simulation_selected_agents = {}

        # outcomes of the simulation
        simulation_outcomes_child = {}
        simulation_outcomes_class = {}

        # select population data
        if(population_name == 'nomination'):
            population = self.PeerNominatedDataPopulation
        elif(population_name == 'communication'):
            population = self.CommunicationDataPopulation

        # set initial PA to snapshot based on pre-intervention simulations
        #print('PRE SNAPSHOT',population.graph.nodes(data='PA'))
        nx.set_node_attributes(population.graph, values=preintervention_snapshot_PA, name='PA')
        #print('POST SNAPSHOT',population.graph.nodes(data='PA'))

        # run the model per intervention
        for intervention in intervention_strategies:
            simulation_outcomes_child[intervention] = {}
            simulation_outcomes_class[intervention] = {}
            simulation_selected_agents[intervention] = {}

            for class_population in population.get_class_graphs(population.graph):
                class_id = list(class_population.nodes(data='class'))[1][1]
                simulation_outcomes_child[intervention][str(class_id)] = {}
                simulation_outcomes_class[intervention][str(class_id)] = {}
                simulation_selected_agents[intervention][str(class_id)] = {}

                cl_pop = class_population.copy()

                # selects influential agents
                influential_agent_selection = population.select_influential_agents(cl_pop, percent, intervention, False)
                # update graph with intervention: increased PA in influential agents
                cl_pop_new = influential_agent_selection[0]
                selected_agents = influential_agent_selection[1]

                #running the simulation, executing the model with every timestamp
                for t in range(0,time):
                    sim_pop = self.model.execute(cl_pop_new,t)
                
                outcomes_in_dict = utils.get_PA_dictionary(sim_pop)
                simulation_outcomes_child[intervention][str(class_id)] = outcomes_in_dict
                simulation_outcomes_class[intervention][str(class_id)] = outcomes_in_dict.mean(axis=1)
                simulation_selected_agents[intervention][str(class_id)] = selected_agents


        # influential agents
        df_agents_list = []
        for outer_dict in simulation_selected_agents.items():
            for intv in outer_dict[1]:
                if(intv != 'nointervention'):
                    df_agents_list.append([outer_dict[0],intv,outer_dict[1][intv]])

        df_agents = pd.DataFrame(df_agents_list, columns = ["Intervention", "SchoolClass", "InfluenceAgents"])

        if writeToExcel:
            self.interventionResultsToExcel(simulation_outcomes_child, intervention_strategies, population_name, run_no)

            directory = '../output/simulation/'+population_name+'/influentialAgents/'
            if not os.path.exists(directory):
                os.makedirs(directory)
            df_agents.to_excel(directory + 'run' + str(run_no) + '.xlsx')


        return simulation_outcomes_child, simulation_outcomes_class, df_agents

            

    def interventionResultsToExcel(self,results,intervention_strategies, population_name, run_no):
        '''
        Saves the simulated intervention outcomes in excel file. 
        
        Args:
            results (dictionary): simulation output
            intervention_strategies (string): interventions
            population_name (string): peer-nomination network or online-communication network
            run_no (int): run number
        '''
        for intervention in intervention_strategies:
            res_intervention = results[intervention]
            directory = '../output/simulation/' + population_name + '/' + intervention

            if not os.path.exists(directory):
                os.makedirs(directory)

            filename = directory + '/run' + str(run_no) + '.xlsx'
            writer = pd.ExcelWriter(filename, engine='xlsxwriter')

            for class_id, res in res_intervention.items():

                #directory='../output/simulation/class'+repr(int(float(class_id)))


                res.to_excel(writer, sheet_name='class'+repr(int(float(class_id))))

            writer.save()

            
    def getSuccessRates(self,results):
        '''
        Calculates the success rates of the simulated interventions.
        
        Args:
            results (dictionary): simulated intervention outcomes
        
        Returns:
            dataframe: dataframe containing success rates per class per simulated interventions
        '''
        success_rates = []
        for class_id,res in results.items():
            for i in self.input_args['intervention_strategy']:
                outcome = res[i]
                success_rates.append([class_id, i, self.get_change(outcome[364],outcome[0]), outcome[0] , outcome[364]])
            
        success_rates = pd.DataFrame(success_rates, columns = ["SchoolClass", "Intervention", "SuccessRate", "StartIntervention", "EndIntervention"])
        
        
        if(self.input_args['writeToExcel'] == True):
            directory='../output/classesSummary'
            if not os.path.exists(directory):
                os.makedirs(directory)
            filename=directory+'/SuccessRates.xls'
            writer = pd.ExcelWriter(filename)
            success_rates.to_excel(writer,'InterventionDif')
            writer.save()
        
        return success_rates

    def get_change(self,current, previous):
        '''
        Helper method for getSuccessRate. Calculates the success rate in percentage.
        
        Args:
            current (float): last time point (end) of the simulation
            previous (float): first time point (start) of the simulation
        
        Returns:
            Integer: percentage of change (success rate)
        '''
        if current == previous:
            return 0
        try:
            return round(((current - previous)/previous)*100.0,2)
        except ZeroDivisionError:
            return 0 

    
    def plot_interventions_per_participant(self,results):    
        '''
         Saves (displays) plots of intervention results per participant per class, based on the dictionary input. 
        
        Args:
            results (dictionary): simulated intervention outcomes
        '''

        for class_id,res in results.items():            

            directory='../output/class'+class_id
            if not os.path.exists(directory):
                os.makedirs(directory)

            document = Document()
            filename=directory+'/'+'class'+class_id+'_Interventions_Detailed_Per_Child.docx'
            document.add_paragraph('All Intervention plots for class ' + class_id + ' per participant')
                
            for i in self.input_args['intervention_strategy']:
                ax = res[i].plot(figsize=((15,10)))
                ax.legend(loc="upper right")
                ax.set_title('Class:'+ class_id+ ' Intervention:'+ i)
                fig =ax.get_figure()
                titlefig=directory +'/'+class_id.strip('\'')+i.strip('\'')+'.png'
                fig.savefig(titlefig)
                document.add_picture(titlefig, width=Inches(7))
                if not self.input_args['save_png']:
                    os.remove(titlefig)

            document.save(filename)
            
            
    def plot_interventions_per_class(self,results_avg):
        '''
         Saves (displays) plots of intervention results per class, based on the dictionary input. 
        
        Args:
            results_avg (dictionary): simulated intervention outcomes (averaged per class)
        '''
        for class_id,res in results_avg.items():
            plt.figure(figsize=((15,10)))
            plt.xlim(0,364)
            plt.xlabel('Days')
            plt.ylabel('Mean PA')
            j = 0
            for i in self.input_args['intervention_strategy']:
                results_avg[class_id][i].plot(color=self.input_args['intervention_color'][j],label= i)
                j = j + 1
            plt.legend(title='All Interventions '+  class_id, loc="upper right")
            
            
    def plot_interventions_averaged(self,results_avg):
        '''
         Saves (displays) plots of average intervention results of all classes
        
        Args:
            results_avg (dictionary): simulated intervention outcomes (averaged per class)
        '''
        all_averaged = {}
        for i in self.input_args['intervention_strategy']:
            temp_res = pd.Series([], dtype = float)
            counter = 0
            for class_id,res in results_avg.items():
                temp_res = temp_res.add(results_avg[class_id][i],fill_value=0)
                counter = counter + 1
            all_averaged[i] = temp_res/counter

        plt.figure(figsize=((15,10)))
        plt.xlim(0,364)
        plt.xlabel('Days')
        plt.ylabel('Mean PA')
        j = 0
        for i in self.input_args['intervention_strategy']:
            all_averaged[i].plot(color=self.input_args['intervention_color'][j],label= i)
            j = j + 1
        plt.legend(title='All Interventions ', loc="upper right")
            
            
    def heatmap(self,success_rates):
        '''
         DEPRECATED Heatmap of success rates
        '''        
        success_rates  = success_rates.groupby(['SchoolClass'])['SuccessRate'].mean().reset_index()
        hm = success_rates[['SchoolClass','SuccessRate']].sort_values('SuccessRate',ascending=False)
        hm=hm.drop(hm.index[len(hm)-1])
        yrows=[1,1,1,1,1,2,2,2,2,2,3,3,3,3,3,4,4,4,4,4,5,5,5,5,5]
        hm=hm.assign(Yrows=yrows)
        xrows=[1,2,3,4,5,1,2,3,4,5,1,2,3,4,5,1,2,3,4,5,1,2,3,4,5]
        hm=hm.assign(Xrows=xrows)
        
        writer = pd.ExcelWriter('../output/heatmap.xlsx')
        hm.to_excel(writer, 'heatmap')
        
        cls=(np.asarray(hm['SchoolClass'])).reshape(5,5)
        perc=(np.asarray(hm['SuccessRate'])).reshape(5,5)
        result=hm.pivot(index='Yrows',columns='Xrows',values='SuccessRate')
        labels=(np.asarray(["Class {0} \n \n {1:.2f}%".format(c,p)
                          for c,p in zip(cls.flatten(),perc.flatten())])
               ).reshape(5,5)
        fig, ax =plt.subplots(figsize=(15,10))
        title= "Success Rates per Class"
        plt.title(title,fontsize=18)
        ttl=ax.title
        ttl.set_position([0.5,1.05])
        ax.set_xticks([])
        ax.set_yticks([])
        ax.axis('off')
        sns.heatmap(result,annot=labels,annot_kws={"size": 16},fmt="",cmap="Blues",linewidths=0.30,ax=ax)
        plt.show()