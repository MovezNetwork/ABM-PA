import networkx as nx
import json
import numpy as np
import os
import pandas as pd
import time
from docx import Document 
from xlwt import Workbook
import collections
from docx.shared import Inches
import matplotlib.pyplot as plt
from src.models import diffuse_behavior_PA, contagion_model, get_graphs_PA_df_detailed
from src.graph import get_empirical,generate_network_PA
from src.interventions import apply_intervention_random_nodes,apply_intervention_pal
from src.interventions import apply_interventions_centrality
from scipy.stats import pearsonr
import seaborn as sns
from itertools import chain

def get_class_graphs(graph, level_f='', writeToFile=False, label='all', c_list=[]):

    '''
    
    Return list of NetworkX subgraphs (each subgraph is a separate class).
    
    Specify wanted classes in c_list, generate all classes by default.
    
    writeToFile - set to True, if you want to generate gephi compatible files. Files are saved on the local path under
                  Subgraphs directory.
                  
    label - the network type of the input graph.
    
    '''
    class_list = []
#     if(writeToFile):
#         directory='output/ClassesSummary/GephiSubgraphs'
#         if not os.path.exists(directory):
#             os.makedirs(directory)
    # if list is empty, we want all the classes
    if not c_list:
        class_list = [67, 71, 72, 74, 77, 78, 79, 81, 83, 86, 100, 101, 103, 121, 122, 125, 126, 127, 129, 130, 131, 133, 135, 136, 138, 139]
    else:
        class_list= c_list
    
    class_dictionary = {}

    for c in class_list:
        class_dictionary[c] = []

    for node, key in graph.nodes.data('class'):
        if key in class_dictionary: 
            class_dictionary[int(key)].append(node)
        
    list_subgraphs = [] 
    cent_dict = {}
    
    for c in class_list:
        subgraph = graph.subgraph(class_dictionary[c]).copy()
        subgraph.graph['networkType']=label
        subgraph.graph['class']=c
        
        if(writeToFile):
                directory='../output/Class'+repr(int(c))
                if not os.path.exists(directory):
                    os.makedirs(directory)
                g_file = directory+'/class_' + repr(c) + '_' + label + '.gexf'
                try:
                    nx.write_gexf(subgraph, g_file)
                except IOError as e:
                    errno, strerror = e.args
                    print("I/O error({0}): {1}".format(errno,strerror))
           
        list_subgraphs.append(subgraph)
        
    return  list_subgraphs

def class_network_details(graph):
    
    '''
        Return two dictionaries suitable for analysis of input graph's network topology:
        
            c2 - network topology characteristics of the input graph (class).
            listChildren - list of child dictionaries, containing information about child characteristics, but also
                           child's role as a node in the network (class).
        
        New network features should be added here (if needed for analysis).
        
        Input - graph: representing a single class
    '''
    
    sg=graph.copy()
    sumPA=0 
    sumBMI=0
    #0=boy ; 1=girl
    totalBoys=0
    totalGirls=0
    totalChildren=sg.number_of_nodes()
    classPA=[]
    classBMI=[]

    c2 = {'ID': '', 'networkType':'', 'numChildren': '', 'numBoys':'', 'numGirls': '', 'ratio': '', 'numConnections': '', 'density':'', 'average_shortest_path_length':'','avg_degree':'','degree_assortativity':'','avgEnv':'','avgPAw1':'','avgPAw2':'','avgPAw3':'','avgPAw4':'','avgBMI':'','classOutDegree':'','classInDegree':'','classCloseness':'','classBetweeness':''}
    listChildren=[]


    dout=dict(sg.out_degree())
    din=dict(sg.in_degree())
    deigen=dict(nx.eigenvector_centrality(sg))
    dcloseness=dict(nx.closeness_centrality(sg))
    dbetweenness=dict(nx.betweenness_centrality(sg))
    
    #centrality measurements on class level
    max_ind=max(din.values())
    dividor_ind=(len(din.values())-1)*(max_ind-1)
    sum_ind=0 
    
    max_outd=max(dout.values())
    dividor_outd=(len(dout.values())-1)*(max_outd-1)
    sum_outd=0   
    
    max_close=max(dcloseness.values())
    dividor_close=(len(dcloseness.values())-1)*(max_close-1) if (len(dcloseness.values())-1)*(max_close-1)!=0 else 1
    sum_close=0
    
    max_beetwn=max(dbetweenness.values())
    dividor_beetwn=(len(dbetweenness.values())-1)*(max_beetwn-1)
    sum_beetwn=0    
    
           #https://networkx.github.io/documentation/stable/reference/algorithms/generated/networkx.algorithms.assortativity.average_neighbor_degree.html#networkx.algorithms.assortativity.average_neighbor_degree
    # Returns the average degree of the neighborhood of each node.
    davgneighdeg=dict(nx.average_neighbor_degree(sg, source='out', target='out',weight='weight'))

    networkType=dict(sg.graph)['networkType']
    avgEnv=0;
    
    # build the average weight dataframe
    avgw=pd.DataFrame(pd.np.empty((0, 2)))
    avgw.columns=['ID','avgW']
    rowCounter=1
    for (u, v, wt) in sg.edges.data('weight'):
        if u in avgw['ID'].values:
            ind=avgw[avgw['ID']==u]['avgW'].index[0]
            val=avgw[avgw['ID']==u]['avgW'].values[0]
            avgw.loc[ind]=[u,wt+val]
        else:
            avgw.loc[rowCounter]=[u,wt]
            rowCounter=rowCounter+1
    
    ind_zero=[]
    # looping through the children nodes
    for nodedata in sg.nodes().data():

        classID=nodedata[1]['class']
        classPA.append(nodedata[1]['PA'])
        classBMI.append(nodedata[1]['bmi'])
        child = {'ID': '', 'class': '','networkType':'', 'PA':'', 'BMI': '', 'env':'', 'gender':'','in_degree':'', 'out_degree':'', 'eigen_centrality':'','closeness_centrality':'','betweenness_centrality':'','average_neighbor_degree':'', 'avg_weight_strength':''}
        child['ID']=nodedata[0]
        child['class']=classID
        child['networkType']=networkType
        child['PA']=nodedata[1]['PA']
        child['BMI']=nodedata[1]['bmi']
        child['env']=nodedata[1]['env']
        child['gender']=nodedata[1]['gender']
        child['in_degree']=din[nodedata[0]]
        sum_ind=sum_ind+(max_ind-din[nodedata[0]])
        if din[nodedata[0]]==0:
            ind_zero.append(nodedata[0])
        child['out_degree']=dout[nodedata[0]]
        sum_outd=sum_outd+(max_outd-dout[nodedata[0]])
        child['eigen_centrality']=deigen[nodedata[0]]
        child['closeness_centrality']=dcloseness[nodedata[0]]
        sum_close=sum_close+(max_close-dcloseness[nodedata[0]])
        child['betweenness_centrality']=dbetweenness[nodedata[0]]
        sum_beetwn=sum_beetwn+(max_beetwn+dbetweenness[nodedata[0]])
        child['average_neighbor_degree']=davgneighdeg[nodedata[0]]
        child['avg_weight_strength']=avgw[avgw['ID']==nodedata[0]]['avgW'].values[0]/dout[nodedata[0]]
        
        avgEnv=avgEnv+nodedata[1]['env']
        
        if(nodedata[1]['gender']==1.0):
            totalGirls=totalGirls+1
        elif(nodedata[1]['gender']==0.0):
            totalBoys=totalBoys+1
            
        listChildren.append(child)
        
    empirical=get_empirical(classes=[classID])
    
    print(classID)
    print('isolated children')
    print(ind_zero)
    print('*****************')
    
    c2['ID']=classID
    c2['numChildren']=totalChildren
    c2['numBoys']=totalBoys
    c2['numGirls']=totalGirls
    c2['ratio']='M:'+str(int((totalBoys/totalChildren)*100))+', F:'+str(int((totalGirls/totalChildren)*100))
    c2['numConnections']=sg.number_of_edges()
    c2['networkType']=networkType
    c2['avgEnv']=round(avgEnv/totalChildren,2)
    c2['avgPAw1']=empirical[1].mean()
    c2['avgPAw2']=empirical[2].mean()
    c2['avgPAw3']=empirical[3].mean()
    c2['avgPAw4']=empirical[4].mean()
    c2['avgBMI']=sumBMI/totalChildren
    c2['density']=round(nx.density(sg),2)
    c2['average_shortest_path_length']=round(nx.average_shortest_path_length(sg),2)
    c2['classOutDegree']=round(sum_outd/dividor_outd,2)
    c2['classInDegree']=round(sum_ind/dividor_ind,2)
    c2['classCloseness']=round(sum_close/dividor_close,2)
    c2['classBetweeness']=round(sum_beetwn/dividor_beetwn,2)
    # Assortativity quantifies the tendency of nodes being connected to similar nodes in a complex network.
    # Tendency for nodes of high degree (resp.low degree) in a graph to be connected to high degree nodes (resp. to low degree ones)
    c2['degree_assortativity']=round(nx.degree_assortativity_coefficient(sg),2)
    
    return c2,listChildren
    

def get_df_class_children_topology_analysis(graphAll=[],graphGen=[],graphFrd=[],network=[], generateGephiFiles=False, generateExcel=False, c_list=[]):
    
    '''
        Return two dataframes suitable for network topology analysis:
            1. classdf - network characteristics of all classes.
            2. childrendf - personal and network characteristics of all children in classes.
            
        Input:
            network - the network type graphs we want to generate.
            c_list - list of classes of interest, all classes by default.
            
    '''
    
    
    # if list is empty, we want all the classes
    class_list = c_list if c_list else [67, 71, 72, 74, 77, 78, 79, 81, 83, 86, 100, 101, 103, 121, 122, 125, 126, 127, 129, 130, 131, 133, 135, 136, 138, 139]   
    networktypes= network if network else ['all', 'gen', 'friend']
    list_subgraphs = []
    
    classCounter=0
    childCounter=0
    for label in networktypes:
        # check if we need to create graphs
        if label=='all':
            #no subgraphs need to create them
            if not graphAll:
                graph = generate_network_PA(level_f='',label=label)
                classes_graph=get_class_graphs(graph.copy(),writeToFile=generateGephiFiles, label=label,c_list=c_list)
                list_subgraphs.extend(classes_graph)
            else:
                classes_graph=graphAll
                list_subgraphs.extend(classes_graph)
        elif label=='gen':
            #no subgraphs need to create them
            if not graphGen:
                graph = generate_network_PA(level_f='',label=label)
                classes_graph=get_class_graphs(graph.copy(),writeToFile=generateGephiFiles, label=label,c_list=c_list)
                list_subgraphs.extend(classes_graph)
            else:
                classes_graph=graphGen
                list_subgraphs.extend(classes_graph)
        elif label=='friend':
            if not graphFrd:
                graph = generate_network_PA(level_f='',label=label)
                classes_graph=get_class_graphs(graph.copy(),writeToFile=generateGephiFiles, label=label,c_list=c_list)
                list_subgraphs.extend(classes_graph)
            else:
                classes_graph=graphFrd
                list_subgraphs.extend(classes_graph)
        
        for sg in classes_graph:
            c2, allChildren=class_network_details(sg.copy())

            #build class dataframe
            if(classCounter==0):
                classdf=pd.DataFrame(list(c2.items())).T
                classdf.columns = classdf.iloc[0]
                classdf=classdf.drop(classdf.index[0])
                classCounter=classCounter+2
            else:
                classdf.loc[classCounter] = list(c2.values())
                classCounter=classCounter+1


            #build children dataframe
            for child in allChildren:

                if(childCounter==0):
                    childdf=pd.DataFrame(list(child.items())).T
                    childdf.columns = childdf.iloc[0]
                    childdf=childdf.drop(childdf.index[0])
                    childCounter=childCounter+2
                else:
                    childdf.loc[childCounter] = list(child.values())
                    childCounter=childCounter+1
          
        if generateExcel:
            directory='../output/ClassesSummary'
            if not os.path.exists(directory):
                os.makedirs(directory)
            filename=directory+'/networkanalysis_'+label+'.xlsx'
            writer = pd.ExcelWriter(filename)
            classdf.to_excel(writer,'Class')
            childdf.to_excel(writer,'Children')
            writer.save()
               
    return list_subgraphs, classdf, childdf

def get_classes_intervention_results():
    
    '''
    
    Return a dictionary containing all the interventions applied per class. The generated dictionary depends on the desired 
    output, as specified by the input parameters:
        network - which network types we want to investigate (all , gen, friend)
        model - what models to use in intervention (diffusion, contagion)
        per - how much percent of each class to target the intervention at ( 10, 15, 20%)
        c_list - which classes are we interested in, all by default  
        generateGephiFiles - should we generate gephi files for all the classes (graphs)
    
    '''
 
    start = time.time()
    
    gr=nx.DiGraph()
    c_list=[]
    
    try:
        input_simulation = json.loads(open('../input/simulation.json').read())
    except Exception as ex:
        print('simulation.json does not exist!')
        print(ex)
        return

        
    # parameters for the diffusion model [threshold, I_PA] 
    parameters_all = input_simulation['parameters_all']
    parameters_gen = input_simulation['parameters_gen']
    parameters_friend = input_simulation['parameters_friend']
    
    # parameters contagion model
    delta=input_simulation['delta']

    networktypes = input_simulation['network'] 
    perc = input_simulation['percent'] 
    model = input_simulation['model'] 
    generateGephiFiles = input_simulation['generateGephiFiles'] 
    writeToExcel = input_simulation['writeToExcel'] 
    class_list = input_simulation['classes']

    # Where the results are going to be saved
    classes_results=[]
    # the current dictionary of interest
    results_dict={}
    classExists=False;
    
    graphAll=[]
    graphGen=[]
    graphFrd=[]   
  
    # create the empty dict with class ID info only
    for classID in class_list:
        #loop the classes dictionaries
        for r_dict in classes_results:
            if(r_dict['class']==classID):
                results_dict=r_dict
                classExists=True
        # class not found, create the class and append to the list of dictionaries
        
        if not classExists:
            r_dict = {'class':'',
            'contagion':{
                    'outdegree':{'all':{10:{}, 15: {}, 20:{}}, 'gen': {10:{}, 15: {}, 20:{}}, 'friend':{10:{}, 15: {}, 20:{}}},
                    'indegree':{'all':{10:{}, 15: {}, 20:{}}, 'gen': {10:{}, 15: {}, 20:{}}, 'friend':{10:{}, 15: {}, 20:{}}},
                    'closeness':{'all':{10:{}, 15: {}, 20:{}}, 'gen': {10:{}, 15: {}, 20:{}}, 'friend':{10:{}, 15: {}, 20:{}}},
                    'betweenness':{'all':{10:{}, 15: {}, 20:{}}, 'gen': {10:{}, 15: {}, 20:{}}, 'friend':{10:{}, 15: {}, 20:{}}},
                    'high_risk': {'all':{10:{}, 15: {}, 20:{}}, 'gen': {10:{}, 15: {}, 20:{}}, 'friend':{10:{}, 15: {}, 20:{}}},
                'maxpal':{'all':{10:{}, 15: {}, 20:{}}, 'gen': {10:{}, 15: {}, 20:{}}, 'friend':{10:{}, 15: {}, 20:{}}},
                'minpal':{'all':{10:{}, 15: {}, 20:{}}, 'gen': {10:{}, 15: {}, 20:{}}, 'friend':{10:{}, 15: {}, 20:{}}},
                    'vulnerability': {'all':{10:{}, 15: {}, 20:{}}, 'gen': {10:{}, 15: {}, 20:{}}, 'friend':{10:{}, 15: {}, 20:{}}},
                    'random': {'all':{10:{}, 15: {}, 20:{}}, 'gen': {10:{}, 15: {}, 20:{}}, 'friend':{10:{}, 15: {}, 20:{}}},
                    'nointervention': {'all':{10:{}, 15: {}, 20:{}}, 'gen': {10:{}, 15: {}, 20:{}}, 'friend':{10:{}, 15: {}, 20:{}}},
                    'optimized': {'all':{10:{}, 15: {}, 20:{}}, 'gen': {10:{}, 15: {}, 20:{}}, 'friend':{10:{}, 15: {}, 20:{}}}
             },
            'diffusion':{
                    'outdegree':{'all':{10:{}, 15: {}, 20:{}}, 'gen': {10:{}, 15: {}, 20:{}}, 'friend':{10:{}, 15: {}, 20:{}}},
                    'indegree':{'all':{10:{}, 15: {}, 20:{}}, 'gen': {10:{}, 15: {}, 20:{}}, 'friend':{10:{}, 15: {}, 20:{}}},
                    'closeness':{'all':{10:{}, 15: {}, 20:{}}, 'gen': {10:{}, 15: {}, 20:{}}, 'friend':{10:{}, 15: {}, 20:{}}},
                    'betweenness':{'all':{10:{}, 15: {}, 20:{}}, 'gen': {10:{}, 15: {}, 20:{}}, 'friend':{10:{}, 15: {}, 20:{}}},
                    'high_risk': {'all':{10:{}, 15: {}, 20:{}}, 'gen': {10:{}, 15: {}, 20:{}}, 'friend':{10:{}, 15: {}, 20:{}}},
                                'maxpal':{'all':{10:{}, 15: {}, 20:{}}, 'gen': {10:{}, 15: {}, 20:{}}, 'friend':{10:{}, 15: {}, 20:{}}},
                'minpal':{'all':{10:{}, 15: {}, 20:{}}, 'gen': {10:{}, 15: {}, 20:{}}, 'friend':{10:{}, 15: {}, 20:{}}},
                    'vulnerability': {'all':{10:{}, 15: {}, 20:{}}, 'gen': {10:{}, 15: {}, 20:{}}, 'friend':{10:{}, 15: {}, 20:{}}},
                    'random': {'all':{10:{}, 15: {}, 20:{}}, 'gen': {10:{}, 15: {}, 20:{}}, 'friend':{10:{}, 15: {}, 20:{}}},
                    'nointervention': {'all':{10:{}, 15: {}, 20:{}}, 'gen': {10:{}, 15: {}, 20:{}}, 'friend':{10:{}, 15: {}, 20:{}}},
                    'optimized': {'all':{10:{}, 15: {}, 20:{}}, 'gen': {10:{}, 15: {}, 20:{}}, 'friend':{10:{}, 15: {}, 20:{}}}
             },
            'empirical': {'all':{}, 'gen': {}, 'friend':{}}
        }
            r_dict['class']=classID
            classes_results.append(r_dict)


    
    dictionaryCreated=False
    
    for label in networktypes:
        # first generate the graph of interest
        graph=gr.copy()
        
        if not graph:
            print('Generating Graph...')
            graph = generate_network_PA(level_f='',label=label)
        # generate all subgraphs - each class is represented with its own graph; we can specify the classes IDs we are interested in by filling the list c_list; c_list=[] means we want all classes
            classes_graph=get_class_graphs(graph.copy(),writeToFile=generateGephiFiles, label=label,c_list=c_list)
        else:
            classes_graph=graph.copy()
            
        if label=='all':
            graphAll=classes_graph.copy()
        elif label=='gen':
            graphGen=classes_graph.copy()
        elif label=='friend':
            graphFrd=classes_graph.copy()
            
        classCounter=0
        numClasses=len(c_list)    
        #per class intervention results
        for subg in classes_graph:
            #get current subgraph class
                classID=list(subg.nodes(data='class'))[1][1]
                subg.graph['int']='for max influence'
                

#             #loop the classes dictionaries
                for r_dict in classes_results:
                    if(r_dict['class']==classID):
                        results_dict=r_dict
                #write the empirical data for the particular network type
                results_dict['empirical'][label] = get_empirical(classes=[classID])
                #replace the modified dictionary, in the same location
#                 cr[cr.index(current_dict)]=d_new
                # loop the selected percentages
                for percent in perc:
            
#                     print('Network Type ' + repr(label)+'; Percent' + repr(percent))
#                     print('Class ID ' + repr(results_dict['class']))
                    classCounter=classCounter+1



                    if(model=='diffusion'):
                        g_no_intervention_dif=diffuse_behavior_PA(graph=g_no_intervention_dif, thres_PA=parameters_all[0], I_PA=parameters_all[1], years=1)
                        results_dict['diffusion']['nointervention'][label][percent]=get_graphs_PA_df_detailed(g_no_intervention_dif)
                    elif(model=='contagion'):
                        g_no_intervention_con=contagion_model(graph=g_no_intervention_con, years=1, delta=delta, model='weighted')
                        results_dict['contagion']['nointervention'][label][percent]=get_graphs_PA_df_detailed(g_no_intervention_con)
                    elif(model=='all'):
                        g_no_intervention_dif=diffuse_behavior_PA(graph=g_no_intervention_dif, thres_PA=parameters_all[0], I_PA=parameters_all[1], years=1)
                        results_dict['diffusion']['nointervention'][label][percent]=get_graphs_PA_df_detailed(g_no_intervention_dif)
                        g_no_intervention_con=contagion_model(graph=g_no_intervention_con, years=1, delta=delta, model='weighted')
                        results_dict['contagion']['nointervention'][label][percent]=get_graphs_PA_df_detailed(g_no_intervention_con)



                    #run out-degree centrality 
                    g_outdegree = apply_interventions_centrality(graph=subg.copy(),perc=percent/100,centrality_type='outdegree')
                    g_outdegree_dif=g_outdegree.copy()
                    g_outdegree_con=g_outdegree.copy()
                    
                    if(model=='diffusion'):
                        g_outdegree_dif=diffuse_behavior_PA(graph=g_outdegree_dif, thres_PA=parameters_all[0], I_PA=parameters_all[1], years=1)
                        results_dict['diffusion']['outdegree'][label][percent]= get_graphs_PA_df_detailed(g_outdegree_dif)
                    elif(model=='contagion'):
                        g_outdegree_con=contagion_model(graph=g_outdegree_con, years=1, delta=delta, model='weighted')
                        results_dict['contagion']['outdegree'][label][percent]=get_graphs_PA_df_detailed(g_outdegree_con)
                    elif(model=='all'):
                        g_outdegree_dif=diffuse_behavior_PA(graph=g_outdegree_dif, thres_PA=parameters_all[0], I_PA=parameters_all[1], years=1)
                        results_dict['diffusion']['outdegree'][label][percent]= get_graphs_PA_df_detailed(g_outdegree_dif)
                        g_outdegree_con=contagion_model(graph=g_outdegree_con, years=1, delta=delta, model='weighted')
                        results_dict['contagion']['outdegree'][label][percent]=get_graphs_PA_df_detailed(g_outdegree_con)

                        
                    #run in-degree centrality 
                    g_indegree = apply_interventions_centrality(graph=subg.copy(),perc=percent/100,centrality_type='indegree')
                    g_indegree_dif=g_indegree.copy()
                    g_indegree_con=g_indegree.copy()
                    
                    if(model=='diffusion'):
                        g_indegree_dif=diffuse_behavior_PA(graph=g_indegree_dif, thres_PA=parameters_all[0], I_PA=parameters_all[1], years=1)
                        results_dict['diffusion']['indegree'][label][percent]= get_graphs_PA_df_detailed(g_indegree_dif)
                    elif(model=='contagion'):
                        g_indegree_con=contagion_model(graph=g_indegree_con, years=1, delta=delta, model='weighted')
                        results_dict['contagion']['indegree'][label][percent]=get_graphs_PA_df_detailed(g_indegree_con)
                    elif(model=='all'):
                        g_indegree_dif=diffuse_behavior_PA(graph=g_indegree_dif, thres_PA=parameters_all[0], I_PA=parameters_all[1], years=1)
                        results_dict['diffusion']['indegree'][label][percent]= get_graphs_PA_df_detailed(g_indegree_dif)
                        g_indegree_con=contagion_model(graph=g_indegree_con, years=1, delta=delta, model='weighted')
                        results_dict['contagion']['indegree'][label][percent]=get_graphs_PA_df_detailed(g_indegree_con)    
                    
                    #run closeness centrality 
                    g_closeness = apply_interventions_centrality(graph=subg.copy(),perc=percent/100,centrality_type='closeness')
                    g_closeness_dif=g_closeness.copy()
                    g_closeness_con=g_closeness.copy()
                    
                    if(model=='diffusion'):
                        g_closeness_dif=diffuse_behavior_PA(graph=g_closeness_dif, thres_PA=parameters_all[0], I_PA=parameters_all[1], years=1)
                        results_dict['diffusion']['closeness'][label][percent]= get_graphs_PA_df_detailed(g_closeness_dif)
                    elif(model=='contagion'):
                        g_closeness_con=contagion_model(graph=g_closeness_con, years=1, delta=delta, model='weighted')
                        results_dict['contagion']['closeness'][label][percent]=get_graphs_PA_df_detailed(g_closeness_con)
                    elif(model=='all'):
                        g_closeness_dif=diffuse_behavior_PA(graph=g_closeness_dif, thres_PA=parameters_all[0], I_PA=parameters_all[1], years=1)
                        results_dict['diffusion']['closeness'][label][percent]= get_graphs_PA_df_detailed(g_closeness_dif)
                        g_closeness_con=contagion_model(graph=g_closeness_con, years=1, delta=delta, model='weighted')
                        results_dict['contagion']['closeness'][label][percent]=get_graphs_PA_df_detailed(g_closeness_con)                        
                        
                    #run betweenness centrality 
                    g_betweenness = apply_interventions_centrality(graph=subg.copy(),perc=percent/100,centrality_type='betweenness')
                    g_betweenness_dif=g_betweenness.copy()
                    g_betweenness_con=g_betweenness.copy()
                    
                    if(model=='diffusion'):
                        g_betweenness_dif=diffuse_behavior_PA(graph=g_betweenness_dif, thres_PA=parameters_all[0], I_PA=parameters_all[1], years=1)
                        results_dict['diffusion']['betweenness'][label][percent]= get_graphs_PA_df_detailed(g_betweenness_dif)
                    elif(model=='contagion'):
                        g_betweenness_con=contagion_model(graph=g_betweenness_con, years=1, delta=delta, model='weighted')
                        results_dict['contagion']['betweenness'][label][percent]=get_graphs_PA_df_detailed(g_betweenness_con)
                    elif(model=='all'):
                        g_betweenness_dif=diffuse_behavior_PA(graph=g_betweenness_dif, thres_PA=parameters_all[0], I_PA=parameters_all[1], years=1)
                        results_dict['diffusion']['betweenness'][label][percent]= get_graphs_PA_df_detailed(g_betweenness_dif)
                        g_betweenness_con=contagion_model(graph=g_betweenness_con, years=1, delta=delta, model='weighted')
                        results_dict['contagion']['betweenness'][label][percent]=get_graphs_PA_df_detailed(g_betweenness_con)

                    #run maxpal  
                    g_maxpal = apply_intervention_pal(graph=subg.copy(),perc=percent/100,criteria='max')
                    g_maxpal_dif=g_maxpal.copy()
                    g_maxpal_con=g_maxpal.copy()
                    
                    if(model=='diffusion'):
                        g_maxpal_dif=diffuse_behavior_PA(graph=g_maxpal_dif, thres_PA=parameters_all[0], I_PA=parameters_all[1], years=1)
                        results_dict['diffusion']['maxpal'][label][percent]= get_graphs_PA_df_detailed(g_maxpal_dif)
                    elif(model=='contagion'):
                        g_maxpal_con=contagion_model(graph=g_maxpal_con, years=1, delta=delta, model='weighted')
                        results_dict['contagion']['maxpal'][label][percent]=get_graphs_PA_df_detailed(g_maxpal_con)
                    elif(model=='all'):
                        g_maxpal_dif=diffuse_behavior_PA(graph=g_maxpal_dif, thres_PA=parameters_all[0], I_PA=parameters_all[1], years=1)
                        results_dict['diffusion']['maxpal'][label][percent]= get_graphs_PA_df_detailed(g_maxpal_dif)
                        g_maxpal_con=contagion_model(graph=g_maxpal_con, years=1, delta=delta, model='weighted')
                        results_dict['contagion']['maxpal'][label][percent]=get_graphs_PA_df_detailed(g_maxpal_con)                        
#                     #run minpal  
                    g_minpal = apply_intervention_pal(graph=subg.copy(),perc=percent/100,criteria='min')
                    g_minpal_dif=g_minpal.copy()
                    g_minpal_con=g_minpal.copy()
                    
                    if(model=='diffusion'):
                        g_minpal_dif=diffuse_behavior_PA(graph=g_minpal_dif, thres_PA=parameters_all[0], I_PA=parameters_all[1], years=1)
                        results_dict['diffusion']['minpal'][label][percent]= get_graphs_PA_df_detailed(g_minpal_dif)
                    elif(model=='contagion'):
                        g_minpal_con=contagion_model(graph=g_minpal_con, years=1, delta=delta, model='weighted')
                        results_dict['contagion']['minpal'][label][percent]=get_graphs_PA_df_detailed(g_minpal_con)
                    elif(model=='all'):
                        g_minpal_dif=diffuse_behavior_PA(graph=g_minpal_dif, thres_PA=parameters_all[0], I_PA=parameters_all[1], years=1)
                        results_dict['diffusion']['minpal'][label][percent]= get_graphs_PA_df_detailed(g_minpal_dif)
                        g_minpal_con=contagion_model(graph=g_minpal_con, years=1, delta=delta, model='weighted')
                        results_dict['contagion']['minpal'][label][percent]=get_graphs_PA_df_detailed(g_minpal_con)                        
#                     #run high-risk
#                     g_high_risk = apply_interventions_high_risk(subg.copy(), percent/100)
#                     g_high_risk.graph['int'] = 'high-risk'
#                     if(model=='diffusion'):
#                         diffuse_behavior_PA(graph=g_high_risk, thres_PA=parameters_gen[0], I_PA=parameters_gen[1], years=1)
#                         results_dict['diffusion']['high_risk'][label][percent]= get_graphs_PA_df_detailed(g_high_risk)
#                     elif(model=='contagion'):
#                         contagion_model(graph=g_high_risk, years=1, delta=delta, model='original')
#                         results_dict['contagion']['high_risk'][label][percent]=get_graphs_PA_df_detailed(g_high_risk)
#                     elif(model=='all'):
#                         g_high_risk_cont = g_high_risk.copy()
#                         diffuse_behavior_PA(graph=g_high_risk, thres_PA=parameters_gen[0], I_PA=parameters_gen[1], years=1)
#                         results_dict['diffusion']['high_risk'][label][percent]= get_graphs_PA_df_detailed(g_high_risk)            
#                         contagion_model(graph=g_high_risk_cont, years=1, delta=delta, model='original')
#                         results_dict['contagion']['high_risk'][label][percent]=get_graphs_PA_df_detailed(g_high_risk_cont)  


#                     #run vilnerability
#                     g_vulnerability = apply_interventions_vulnerability(subg.copy(), percent/100)
#                     g_vulnerability_dif=g_vulnerability.copy()
#                     g_vulnerability_con=g_vulnerability.copy()
#                     if(model=='diffusion'):
#                         g_vulnerability_dif=diffuse_behavior_PA(graph=g_vulnerability_dif, thres_PA=parameters_all[0], I_PA=parameters_all[1], years=1)
#                         results_dict['diffusion']['vulnerability'][label][percent]= get_graphs_PA_df_detailed(g_vulnerability_dif)
#                     elif(model=='contagion'):
#                         g_vulnerability_con=contagion_model(graph=g_vulnerability_con, years=1, delta=delta, model='weighted')
#                         results_dict['contagion']['vulnerability'][label][percent]=get_graphs_PA_df_detailed(g_vulnerability_con)
#                     elif(model=='all'):
                        
#                         g_vulnerability_dif=diffuse_behavior_PA(graph=g_vulnerability_dif, thres_PA=parameters_gen[0], I_PA=parameters_gen[1], years=1)
#                         results_dict['diffusion']['vulnerability'][label][percent] = get_graphs_PA_df_detailed(g_vulnerability_dif)            
#                         g_vulnerability_con=contagion_model(graph=g_vulnerability_con, years=1, delta=delta, model='weighted')
#                         results_dict['contagion']['vulnerability'][label][percent]=get_graphs_PA_df_detailed(g_vulnerability_con)

#                     #run max influence
#                     if(model=='diffusion'):
#                         g_diffusion_max_influence = apply_intervention_max_influence(subg.copy(), percent/100, modeltype='diffusion')
#                         g_diffusion_max_influence.graph['int'] = 'maxinfluence'
#                         diffuse_behavior_PA(graph=g_diffusion_max_influence, thres_PA=parameters_all[0], I_PA=parameters_all[1], years=1)
#                         results_dict['diffusion']['optimized'][label][percent]= get_graphs_PA_df_detailed(g_diffusion_max_influence)
#                     elif(model=='contagion'): 
#                         g_contagion_max_influence = apply_intervention_max_influence(subg.copy(), percent/100, modeltype='contagion',delta=delta)
#                         g_contagion_max_influence.graph['int'] = 'maxinfluence'
#                         contagion_model(graph=g_contagion_max_influence, years=1, delta=delta, model='weighted')
#                         results_dict['contagion']['optimized'][label][percent]= get_graphs_PA_df_detailed(g_contagion_max_influence)
#                     elif(model=='all'):
#                         g_diffusion_max_influence = apply_intervention_max_influence(subg.copy(), percent/100, modeltype='diffusion')
#                         g_diffusion_max_influence.graph['int'] = 'maxinfluence'
#                         diffuse_behavior_PA(graph=g_diffusion_max_influence, thres_PA=parameters_all[0], I_PA=parameters_all[1], years=1)
#                         results_dict['diffusion']['optimized'][label][percent]= get_graphs_PA_df_detailed(g_diffusion_max_influence)
#                         g_contagion_max_influence = apply_intervention_max_influence(subg.copy(), percent/100, modeltype='contagion',delta=delta)
#                         g_contagion_max_influence.graph['int'] = 'maxinfluence'
#                         contagion_model(graph=g_contagion_max_influence, years=1, delta=delta, model='weighted')
#                         results_dict['contagion']['optimized'][label][percent]= get_graphs_PA_df_detailed(g_contagion_max_influence)

                    #run random
                    if(model=='diffusion'):
                        g_diffusion_random_list = []
                        for i in range(100):
                            print("Random {0}".format(i), end="\r")
                            g_diffusion_random = apply_intervention_random_nodes(subg.copy())
                            diffuse_behavior_PA(graph=g_diffusion_random, thres_PA=parameters_all[0], I_PA=parameters_all[1], years=1)
                            g_diffusion_random_list.append(get_graphs_PA_df_detailed(g_diffusion_random))
                        # once you get the list of 100 random runningsm concat it to a single dataframe
                        res=pd.concat(g_diffusion_random_list)
                        # this will generate a dataframe of all 100 dataframes so dimensioally it will be 100x365 columns
                        # group by the index, which is a single day, and get the mean of it.
                        res=res.groupby(res.index).mean()
                        results_dict['diffusion']['random'][label][percent]= res                
                    elif(model=='contagion'): 
                        g_contagion_random_list = []
                        for i in range(100):
                            print("Random {0}".format(i), end="\r")
                            g_contagion_random = apply_intervention_random_nodes(subg.copy())
                            contagion_model(graph=g_contagion_random, years=1, delta=delta, model='original')
                            g_contagion_random_list.append(get_graphs_PA_df_detailed(g_contagion_random))   
                        res=pd.concat(g_contagion_random_list)
                        res=res.groupby(res.index).mean()
                        results_dict['contagion']['random'][label][percent]= res
                    elif(model=='all'):
                        g_diffusion_random_list = []
                        for i in range(100):
                            print("Random {0}".format(i), end="\r")
                            g_diffusion_random = apply_intervention_random_nodes(subg.copy())
                            diffuse_behavior_PA(graph=g_diffusion_random, thres_PA=parameters_all[0], I_PA=parameters_all[1], years=1)
                            g_diffusion_random_list.append(get_graphs_PA_df_detailed(g_diffusion_random))
                        res=pd.concat(g_diffusion_random_list)
                        res=res.groupby(res.index).mean()
                        results_dict['diffusion']['random'][label][percent]= res
                        g_contagion_random_list = []
                        for i in range(100):
                            print("Random {0}".format(i), end="\r")
                            g_contagion_random = apply_intervention_random_nodes(subg.copy())
                            contagion_model(graph=g_contagion_random, years=1, delta=delta, model='original')
                            g_contagion_random_list.append(get_graphs_PA_df_detailed(g_contagion_random))   
                        res=pd.concat(g_contagion_random_list)
                        res=res.groupby(res.index).mean()
                        results_dict['contagion']['random'][label][percent]= res
                    
                    classes_results[classes_results.index(results_dict)]=results_dict

    end = time.time()
    print(end - start)
    
    if writeToExcel:
        writeClassesInterventionToExcel(classes_results=classes_results)
    
    return classes_results,graphAll,graphGen,graphFrd


    

def get_interventions_differences(class_dict={}, model=[], label=[], percent=[], intervention=[],writeToExcel=False):
    
    '''
    
    Return a dataframe containing intervention' outcomes results: the differences between day 364 and day 0, of running the simulations for all defined selection strategies. Also the percentage of intervention success rate. All this information is obtained per class. Input parameters
        class_dict - dictionary containing the interventions' results applied per class
        model - computational model used in intervention (diffusion, contagion)
        per - how much percent of each class to target the intervention at ( 10, 15, 20%)
        intervention - selection strategy of interest 
        writeToExcel - Creates an excel with the intervention outcomes' details
    
    '''
    models = model if model else ['diffusion', 'contagion']
    networktypes= label if label else ['all', 'gen', 'friend']
    perc= percent if percent else [10, 15, 20]
    interventions = intervention if intervention else ['optimized', 'outdegree', 'indegree', 'betweenness', 'closeness', 'high_risk', 'maxpal','minpal','vulnerability', 'random', 'nointervention'] 
    
    cdict=class_dict

    result_class=[]
    writeClass=False
    
    for c in cdict:
        for m in models:
            for n in networktypes:
                for p in perc:
                    writeClass=False
#                     print('*** | '+ repr(m)+' model  | '+ repr(n) + ' | ' +  repr(p) + ' %| ***')
                    # get the nointervention data

                    if isinstance(c[m]['nointervention'][n][p], pd.DataFrame):
                        if c[m]['nointervention'][n][p].empty:
                            print('Empty Dictionary '+ repr(m) + ' nointervention '+ repr(n)+ ' ' + repr(p)) 
                        else:
                            
                            no_interv=c[m]['nointervention'][n][p].mean(axis=1)    
                            c2 = {'ID': '', 'networkType':'', 'percent': '', 'model':'', 'optimized': '', 'd0_optimized': '','d364_optimized': '','outdegree': '', 'd0_outdegree': '', 'd364_outdegree': '','indegree': '', 'd0_indegree': '', 'd364_indegree': '', 'betweenness': '', 'd0_betweenness': '', 'd364_betweenness': '', 'closeness': '', 'd0_closeness': '', 'd364_closeness': '', 'maxpal': '', 'd0_maxpal': '', 'd364_maxpal': '','minpal': '', 'd0_minpal': '', 'd364_minpal': '', 'high_risk':'','d0_high_risk':'','d364_high_risk':'','vulnerability':'','d0_vulnerability':'','d364_vulnerability':'','random':'','d0_random':'','d364_random':'','nointervention':'','d0_nointervention':'','d364_nointervention':'','s_optimized':'','e_optimized':'','s_outdegree':'','e_outdegree':'','s_indegree':'','e_indegree':'','s_betweenness':'','e_betweenness':'','s_closeness':'','e_closeness':'','s_maxpal':'','e_maxpal':'','s_minpal':'','e_minpal':'','s_high_risk':'','e_high_risk':'','s_vulnerability':'','e_vulnerability':'','s_random':'','e_random':''}
                            c2['noint']=get_change(no_interv[364],no_interv[0])
                            c2['networkType']=n
                            c2['percent']=p
                            c2['model']=m
                            c2['ID']=c['class']
                            writeClass=True
                            
                            
                    for i in interventions:
                        if isinstance(c[m][i][n][p], pd.DataFrame):
                            if  c[m][i][n][p].empty:
                                print('Empty Dictionary '+ repr(m) + ' '+ repr(i) + ' '+ repr(n)+ ' ' + repr(p))
                            else:  
                                
                                
                                
                                dfint=c[m][i][n][p].mean(axis=1)
                                start='s_'+i.strip('\'')
                                end='e_'+i.strip('\'')
#                                 val='val_'+i.strip('\'')
                                percc='perc_'+i.strip('\'')
                                d0='d0_'+i.strip('\'')
                                d364='d364_'+i.strip('\'')
                                # print(repr(i)+'_start '+ repr(round(start, 4)) + '    '+repr(i)+'_END '+ repr(round(end, 4)))
#                                 c2[i]=get_change(dfint[364],dfint[0])
#                                 c2[start]=get_change(dfint[0],no_interv[0])
#                                 c2[end]=get_change(dfint[364],no_interv[364])
                                c2[i]=dfint[364]-dfint[0]
                                c2[percc]=get_change(dfint[364],dfint[0])
#                                 c2[start]=dfint[0]-no_interv[0]
#                                 c2[end]=dfint[364]-no_interv[364]
                                if i=='nointervention':
                                    c2[d0]=no_interv[0]
                                    c2[d364]=no_interv[346]
                                else:    
                                    c2[d0]=dfint[0]
                                    c2[d364]=dfint[346]
                    
                    if writeClass:
                        result_class.append(c2)
    
    # generate the dataframe
    classCounter=0
    for c2 in result_class:
        if(classCounter==0):
            classdf=pd.DataFrame(list(c2.items())).T
            classdf.columns = classdf.iloc[0]
            classdf=classdf.drop(classdf.index[0])
            classCounter=classCounter+2
        else:
            classdf.loc[classCounter] = list(c2.values())
            classCounter=classCounter+1
    
    if writeToExcel==True:
        directory='../output/ClassesSummary'
        if not os.path.exists(directory):
            os.makedirs(directory)
        filename=directory+'/InterventionDifferences.xls'
        writer = pd.ExcelWriter(filename)
        classdf.to_excel(writer,'InterventionDif')
        writer.save()
                    
    success_rates=classdf[['ID','perc_indegree','perc_betweenness','perc_closeness','perc_random','perc_maxpal','perc_minpal','perc_nointervention']]
    # iloc[:, [1,2,3,4,5,6]] -> depends on what you want to have average
    success_rates['perc_sni']=success_rates.iloc[:, [1,2,3]].mean(axis=1)
    success_rates=success_rates.sort_values('perc_sni')
#     success_rates.loc['avg'] = success_rates.mean()
    success_rates        
            
    return classdf, success_rates
                                
    
def get_intervention_per_child_plots(classes_results=[],save_png=False, create_doc=False, model=[], label=[], percent=[], intervention=[]):    

    '''
    
    Saves (Displays) plots of intervention results per children per class, based on the dictionary input. 
        classes_results - dictionary containing the interventions' results applied per class
        save_png - should the created plots be saved as png files
        create_doc - should you save the png in a single docx
        model - computational model used in intervention (diffusion, contagion)
        percent - how much percent of each class to target the intervention at ( 10, 15, 20%)
    
    '''
    
    for c in classes_results:
        class1=c
        
        if create_doc:
            directory='../output/Class'+repr(int(class1['class']))
            if not os.path.exists(directory):
                os.makedirs(directory)
            document = Document()
            filename=directory+'/'+'Class'+repr(int(class1['class']))+'_Interventions_Detailed_Per_Child.docx'
            document.add_paragraph('All Intervention plots for class ' + repr(int(class1['class'])) + ' with all children')

        models = model if model else ['diffusion', 'contagion','empirical']
        networktypes= label if label else ['all', 'gen', 'friend']
        perc= percent if percent else [10, 15, 20]
        interventions = intervention if intervention else ['optimized', 'outdegree', 'indegree', 'betweenness', 'closeness', 'high_risk', 'maxpal','minpal', 'vulnerability', 'random', 'nointervention']    
        

        print('Plotting data graphs for Class with ID '+  repr(int(class1['class'])))

        for m in models:
            for n in networktypes:
                if(m !='empirical'):
                    for p in perc:
                        for i in interventions:
                            if isinstance(class1[m][i][n][p], pd.DataFrame):
                                if not class1[m][i][n][p].empty:
                                    ax = class1[m][i][n][p].plot(figsize=((15,10)))
                                    ax.legend(loc="upper right")
                                    ax.set_title('Model_Type:'+ repr(m)+ '  Network_Type:' + repr(n) + '  Target' + repr(p)+'%  Intervention:'+ repr(i))
                                    if create_doc:
                                        fig =ax.get_figure()
                                        titlefig=directory +'/'+n.strip('\'')+'_'+m.strip('\'')+'_'+i.strip('\'')+'_'+repr(p)+'.png'
                                        fig.savefig(titlefig)
                                        document.add_picture(titlefig, width=Inches(7))
                                        if not save_png:
                                            os.remove(titlefig)

                                else:
                                    print('Class dataframe empty: ' +repr(m)+repr(i)+repr(n)+repr(p))
                elif(m =='empirical'):
                    if isinstance(class1['empirical'][n], pd.DataFrame):
                            if not class1['empirical'][n].empty:
                                ax = class1['empirical'][n].T.plot(figsize=((15,10)))
                                ax.legend(loc="upper right")
                                ax.set_title('Model_Type:'+ repr(m)+ '  Network_Type:' + repr(n))

                                if create_doc:
                                    fig =ax.get_figure()
                                    titlefig=directory +'/'+n.strip('\'')+'_'+m.strip('\'')+'.png'
                                    fig.savefig(titlefig)
                                    document.add_picture(titlefig, width=Inches(7))
                                    if not save_png:
                                        os.remove(titlefig)
        if create_doc:    
            document.save(filename)

        
def get_intervention_model_comparison_plots(classes_results=[],save_png=False, create_doc=False,model=[], label=[], percent=[], intervention=[]):
    
    '''
    
    Model comparison plots.
    
    This is outdated - not containing the journal version intervention strategies
    
    '''
    
    for c in classes_results:
        class1=c

        if create_doc:
            directory='../output/Class'+repr(int(class1['class']))
            if not os.path.exists(directory):
                os.makedirs(directory)
            document = Document()
            filename=directory+'/'+'Class'+repr(int(class1['class']))+'_Model_Comparisons.docx'
            document.add_paragraph('All Intervention plots for class ' + repr(int(class1['class'])) + ' with all children')

        models = model if model else ['diffusion', 'contagion','empirical']
        networktypes= label if label else ['all', 'gen', 'friend']
        perc= percent if percent else [10, 15, 20]
        interventions = intervention if intervention else ['optimized', 'centrality', 'high_risk', 'vulnerability', 'random', 'nointervention']    

        print('Comparing Models | Class with ID '+  repr(int(class1['class'])))

        for n in networktypes:
                    for p in perc:
                        for i in interventions:
                            if(i !='nointervention'):
                                if isinstance(class1['diffusion'][i][n][p], pd.DataFrame):
                                    
                                    if not class1['diffusion'][i][n][p].empty:
                                        plt.figure(figsize=((15,10)))
                                        plt.title('Diffusion vs Contagion' +' Network_Type:' + repr(n) + '  Target' + repr(p)+'%  Intervention:'+ repr(i))
                                        class1['diffusion'][i][n][p].mean(axis=1).plot(color='red',label='Difusion '+repr(i))
                                        class1['contagion'][i][n][p].mean(axis=1).plot(color='blue',label='Contagion '+repr(i))
                                        class1['diffusion']['nointervention'][n][p].mean(axis=1).plot(color='red',style='-.',label='Difusion No Intervention')
                                        class1['contagion']['nointervention'][n][p].mean(axis=1).plot(color='blue',style='-.',label='Contagion No Intervention')
                                        plt.xlim(0,364)
                                        plt.xlabel('Days')
                                        plt.ylabel('Mean PA')
                                        plt.legend(title=repr(p)+'%', loc="upper right")

                                        if create_doc:
                                            titlefig=directory +'/diffusion_contagion_'+i.strip('\'')+'_'+i.strip('\'')+'_'+repr(p)+'.png'
                                            plt.savefig(titlefig)
                                            document.add_picture(titlefig, width=Inches(7))
                                            if not save_png:
                                                os.remove(titlefig)

                                    else:
                                        print('Class dataframe empty: '+repr(i)+repr(n)+repr(p))


        if create_doc:    
            document.save(filename)


def get_all_interventions_per_model_plots(classes_results=[],save_png=False, create_doc=False,model=[], label=[], percent=[], intervention=[]):
    
    for c in classes_results:
        class1=c
    
        if create_doc:
            directory='../output/Class'+repr(int(class1['class']))
            if not os.path.exists(directory):
                os.makedirs(directory)
            document = Document()
            filename=directory+'/'+'Class'+repr(int(class1['class']))+'_Models.docx'
            document.add_paragraph('Per Model All Intervention | Class ' + repr(int(class1['class'])))

        models = model if model else ['diffusion', 'contagion']
        networktypes= label if label else ['all', 'gen', 'friend']
        perc= percent if percent else [10, 15, 20]
        interventions = intervention if intervention else ['optimized', 'outdegree', 'indegree', 'betweenness', 'closeness', 'high_risk', 'maxpal','minpal', 'vulnerability', 'random', 'nointervention']    

        print('Comparing Models | Class with ID '+  repr(int(class1['class'])))


        createPlot=False

        for m in models:
            for n in networktypes:
                if(m !='empirical'):
                    for p in perc:
                        
                        plt.figure(figsize=((15,10)))
                        plt.xlim(0,364)
                        plt.xlabel('Days')
                        plt.ylabel('Mean PA')
                        
                        if isinstance(class1[m]['optimized'][n][p], pd.DataFrame) and not class1[m]['optimized'][n][p].empty:

                            class1[m]['optimized'][n][p].mean(axis=1).plot(color='red',label=repr(m)+'optimized')
                            createPlot=True
                        if isinstance(class1[m]['outdegree'][n][p], pd.DataFrame) and not class1[m]['outdegree'][n][p].empty:    
                            class1[m]['outdegree'][n][p].mean(axis=1).plot(color='blue',label=repr(m)+'outdegree')
                            createPlot=True
                            
                        if isinstance(class1[m]['indegree'][n][p], pd.DataFrame) and not class1[m]['indegree'][n][p].empty:    
                            class1[m]['indegree'][n][p].mean(axis=1).plot(color='green',label=repr(m)+'indegree')
                            createPlot=True  
                            
                        if isinstance(class1[m]['betweenness'][n][p], pd.DataFrame) and not class1[m]['betweenness'][n][p].empty:    
                            class1[m]['betweenness'][n][p].mean(axis=1).plot(color='orange',label=repr(m)+'betweenness')
                            createPlot=True 
 
                        if isinstance(class1[m]['closeness'][n][p], pd.DataFrame) and not class1[m]['closeness'][n][p].empty:    
                            class1[m]['closeness'][n][p].mean(axis=1).plot(color='purple',label=repr(m)+'closeness')
                            createPlot=True  

                        if isinstance(class1[m]['high_risk'][n][p], pd.DataFrame) and not class1[m]['high_risk'][n][p].empty:
                            

                            class1[m]['high_risk'][n][p].mean(axis=1).plot(color='green',label=repr(m)+'high_risk')
                            createPlot=True
                        if isinstance(class1[m]['maxpal'][n][p], pd.DataFrame) and not class1[m]['maxpal'][n][p].empty:
                            

                            class1[m]['maxpal'][n][p].mean(axis=1).plot(color='pink',label=repr(m)+'maxpal')
                            createPlot=True
                        if isinstance(class1[m]['minpal'][n][p], pd.DataFrame) and not class1[m]['minpal'][n][p].empty:
                            

                            class1[m]['minpal'][n][p].mean(axis=1).plot(color='olive',label=repr(m)+'minpal')
                            createPlot=True                            
                        if isinstance(class1[m]['vulnerability'][n][p], pd.DataFrame) and not class1[m]['vulnerability'][n][p].empty:

                            class1[m]['vulnerability'][n][p].mean(axis=1).plot(color='yellow',label=repr(m)+'vulnerability')
                            createPlot=True
                        if isinstance(class1[m]['nointervention'][n][p], pd.DataFrame) and not class1[m]['nointervention'][n][p].empty:

                            class1[m]['nointervention'][n][p].mean(axis=1).plot(style='-', linewidth=5, color='black',label=repr(m)+'nointervention')
                            createPlot=True

                        if create_doc: 
                            if createPlot:
                                plt.legend(title='All Interventions '+  repr(int(class1['class']))+'_' + n.strip('\'')+ ' '+m.strip('\'') + '  Target' + repr(p)+'%', loc="upper right")
                                createPlot=False
                                titlefig=directory +'/All_Interventions_'+  repr(int(class1['class']))+'_'+n.strip('\'')+'_'+m.strip('\'')+'_'+repr(p)+'.png'
                                plt.savefig(titlefig)
                                document.add_picture(titlefig, width=Inches(7))
                                if not save_png:
                                    os.remove(titlefig)

        if create_doc:    
            document.save(filename)


def get_classes_intervention_comparison_plots(classes_results=[],save_png=False, create_doc=False,model=[], label=[], percent=[], intervention=[]):
    
        
    '''
    
    Visualizes the comparison between classes intervention success, looking at each intervention stategy separately. Saves the result as an doc file and png images.
    
    '''
    
    if create_doc:
        directory='../output/ClassesSummary'
        if not os.path.exists(directory):
            os.makedirs(directory)
        document = Document()
        filename=directory+'/'+'Classes_Summary.docx'
     

    models = model if model else ['diffusion', 'contagion']
    networktypes= label if label else ['all', 'gen', 'friend']
    perc= percent if percent else [10, 15, 20]
    interventions = intervention if intervention else ['optimized', 'outdegree', 'indegree', 'betweenness', 'closeness', 'high_risk', 'maxpal','minpal', 'vulnerability', 'random', 'nointervention']
    
    classCounter=0
    
    for m in models:
        for n in networktypes:
                for p in perc:
                    for i in interventions:
                        classCounter=0
                        sns.set()
                        sns.set_style("darkgrid")
                        sns.set_context("talk", rc={"lines.linewidth": 2.5})
#                         sns.set_style("darkgrid")
                        plt.figure(figsize=((16,16)))               
#                         print('*** | '+ repr(m)+'  | '+ repr(n) + '  | '+ repr(i) + ' | ' +  repr(p) + ' %| ***')
                        for c in classes_results:
                            if isinstance(c[m][i][n][p], pd.DataFrame) and not c[m][i][n][p].empty:
                                
                                    classCounter=classCounter+1

                                    if classCounter==1:
    #                                     print('****create subplot1****') 
                                        ax=plt.subplot(221)
                                        text=""
                                        if i=="high_risk":
                                            text='High Risk Intervention Simulation'
                                        elif i=="outdegree":
                                            text='Out Degree Intervention Simulation'
                                        elif i=="indegree":
                                            text='In Degree Intervention Simulation'
                                        elif i=="betweenness":
                                            text='Betweenness Intervention Simulation'
                                        elif i=="maxpal":
                                            text='Max PAL Intervention Simulation'
                                        elif i=="minpal":
                                            text='Min PAL Intervention Simulation'
                                        elif i=="closeness":
                                            text='Closeness Intervention Simulation'
                                        elif i=="vulnerability":
                                            text='Vulnerability Intervention Simulation'
                                        elif i=="nointervention":
                                            text='No Intervention Simulation' 
                                        ax.set_title(text)
                                        plt.xlim(0,364)
                                        plt.ylim(1.1,2.2)
                                        plt.xlabel('Days')
                                        plt.ylabel('Mean PAL')

                                    c[m][i][n][p].mean(axis=1).plot(label='Class '+repr(c['class']))
    #                                 print('add class'+repr(c['class'])+' to subplot')

                                    if classCounter==6:    
                                        plt.legend(title='',loc=4)
                                        print('create legend for subplot1')
                                        print('****create subplot2****')
#                                         ax=plt.subplot(222)
#                                         ax.set_title("Title for first plot")
                                        plt.xlim(0,364)
                                        plt.ylim(1,2.3)
                                        plt.xlabel('Days')
                                        plt.ylabel('Mean PAL')

                                    if classCounter==7:
                                        print('****create subplot2****')
                                        plt.subplot(222)
                                        plt.xlim(0,364)
                                        plt.ylim(1,2.3)
                                        plt.xlabel('Days')
                                        plt.ylabel('Mean PAL')


                                    if classCounter==13:    
                                        plt.legend(title='', loc=4)
    #                                     print('create legend for subplot2')
    #                                     print('****create subplot3****')
                                        plt.subplot(223)
                                        plt.xlim(0,364)
                                        plt.ylim(1,2.3)
                                        plt.xlabel('Days')
                                        plt.ylabel('Mean PAL')

    #                                 if classCounter==13:


                                    if classCounter==20:    
    #                                     print('create legend for subplot3')
                                        plt.legend(title='', loc=4)
    #                                     print('****create subplot4****')
                                        plt.subplot(224)
                                        plt.xlim(0,364)
                                        plt.ylim(1,2.3)
                                        plt.xlabel('Days')
                                        plt.ylabel('Mean PAL')

    #                                 if classCounter==19:


                                    if classCounter==26: 
    #                                     print('create legend for subplot4')
                                        plt.legend(title='', loc=4)

                                
                        if create_doc:
#                                 print('*** GENERATED: | '+ repr(m)+'  | '+ repr(n) + '  | '+ repr(i) + ' | ' +  repr(p) + ' %| ***')
                                titlefig=directory +'/AllClasses_'+n.strip('\'')+'_'+m.strip('\'')+'_'+ i.strip('\'') + '_'+repr(p)+'.png'
                                plt.savefig(titlefig)
                                document.add_picture(titlefig, width=Inches(7))
                                if not save_png:
                                    os.remove(titlefig)
   
    if create_doc:    
        document.save(filename)
        
def writeClassesInterventionToExcel(classes_results=[]):
    models = ['diffusion', 'contagion', 'empirical']
    networktypes=  ['all', 'gen', 'friend']
    perc= [10, 15, 20]
    interventions = ['optimized', 'outdegree', 'indegree', 'betweenness', 'closeness', 'high_risk','maxpal','minpal', 'vulnerability', 'random', 'nointervention']
    test=classes_results

 


    # loop the classes
    for t in test:
        # add new excel file for each class
        w = Workbook()
        directory='../output/Class'+repr(int(t['class']))
        if not os.path.exists(directory):
            os.makedirs(directory)
        filename=directory+'/'+repr(t['class'])+'.xls'
        for m in models:
            if m!='empirical':
                for i in interventions:
                    for n in networktypes:
                        for p in perc:
                            if isinstance(t[m][i][n][p], pd.DataFrame) and not t[m][i][n][p].empty:
                                # new sheet for each case
                                sheet=m.strip('\'')+i.strip('\'')+n.strip('\'')+repr(p).strip('\'')
                                ws = w.add_sheet(sheet)
                                #loop the dataframe
                                rowNum=0
                                colNum=0
                                # list of column labels in the dataframe = childIDs
                                col=list(t[m][i][n][p])
                                colLen=len(col)
                                #writing the labels in excel sheet
                                ws.write(rowNum,colNum,'Days')      
                                for c in col:
                                    colNum=colNum+1
                                    ws.write(rowNum,colNum,c)


                                for row in t[m][i][n][p].itertuples():
                                    rowNum=rowNum+1
                                    #write the day
                                    ws.write(rowNum,0,row.Index)
                                    #write the values
                                    for c in range(1, colLen+1):
                                        ws.write(rowNum,c,row[c])
            else:
                for nn in networktypes:
                     if isinstance(t['empirical'][nn], pd.DataFrame) and not t['empirical'][nn].empty:
                        # new sheet for each case
                        sheet=m.strip('\'')+nn.strip('\'')
                        ws = w.add_sheet(sheet)
                        emp=t[m][nn].T
                        #loop the dataframe
                        rowNum=0
                        colNum=0
                        # list of column labels in the dataframe = childIDs
                        col=list(emp)
                        colLen=len(col)
                        #writing the labels in excel sheet
                        ws.write(rowNum,colNum,'Wave')      
                        for c in col:
                            colNum=colNum+1
                            ws.write(rowNum,colNum,c)


                        for row in emp.itertuples():
                            rowNum=rowNum+1
                            #write the day
                            ws.write(rowNum,0,row.Index)
                            #write the values
                            for c in range(1, colLen+1):
                                ws.write(rowNum,c,row[c])

        w.save(filename)

def allChildrenInClass():
    pp = pd.read_csv('../data/pp.csv', sep=';', header=0)
    pp['cl'] = pp.Class_Y1
    pp['child'] = pp.Child_Bosse
    # Fill the missing data at Class column with the data from Y1.
    pp['cl'].fillna(pp.Class_Y2, inplace=True)
    pp.index = pp.Child_Bosse
    class_df = pp[['cl','child']]
    class_list=class_df['cl'].tolist()
    
    wanted=[67, 71, 72, 74, 77, 78, 79, 81, 83, 86, 100, 101, 103, 121, 122, 125, 126, 127, 129, 130, 131, 133, 135, 136, 138, 139]
    class_df=class_df[class_df.cl.isin(wanted)]
    allChildrenClass = pd.DataFrame(columns=['Class','NumChildren'])
    for u in wanted:
        occurance=class_list.count(u)
        allChildrenClass.loc[len(allChildrenClass)] = [u,occurance]
        allChildrenClass.to_csv('classAllChildren.csv')
     
    return class_df
        
def get_change(current, previous):
    if current == previous:
        return 0
    try:
        return round(((current - previous)/previous)*100.0,2)
    except ZeroDivisionError:
        return 0  
    
def df_to_excel(df,filename='',sheet=''):             
    writer = pd.ExcelWriter(filename)
    df.to_excel(writer,sheet)
    writer.save()
    
def get_degree_histogram(graph):
    # degree histogram
    degree_sequence = sorted([d for n, d in graph.out_degree()], reverse=True)
    degreeCount = collections.Counter(degree_sequence)
    deg, cnt = zip(*degreeCount.items())
    fig, ax = plt.subplots()
    plt.bar(deg, cnt, width=0.80, color='b')
    plt.title("Degree Histogram")
    plt.ylabel("Count")
    plt.xlabel("Degree")
    ax.set_xticks([d + 0.4 for d in deg])
    ax.set_xticklabels(deg)
    plt.show()       

    
def calculate_pvalues(df):
    df = df.dropna()._get_numeric_data()
    dfcols = pd.DataFrame(columns=df.columns)
    pvalues = dfcols.transpose().join(dfcols, how='outer')
    for r in df.columns:
        for c in df.columns:
            pvalues[r][c] = round(pearsonr(df[r], df[c])[1], 4)
    return pvalues    
    
def get_correlations(dct={},model='all'):
    df_interventions=dct
    net_analysis=pd.ExcelFile('../output/ClassesSummary/networkanalysis_gen.xlsx',sheet_name='Class')
    net_analysis=net_analysis.parse('Class')
    net_analysis=net_analysis[['ID','density','ROutDegreeCentralization','RClosenessCentralization']]
    net_analysis=net_analysis.set_index('ID')
    
    # get the intervention differentes between running the simulation in day0 and day364; day364-day0
    int_diff=get_interventions_differences(class_dict=df_interventions,label=['gen'],percent=[20],intervention=['centrality', 'vulnerability', 'nointervention','high_risk'])
    int_diff=int_diff[['ID', 'model', 'centrality','vulnerability', 'noint','high_risk']]
    int_diff_contagion=int_diff[int_diff.model=='contagion']
    int_diff_diffusion=int_diff[int_diff.model=='diffusion']
    int_diff_contagion=int_diff_contagion.set_index('ID')
    int_diff_diffusion=int_diff_diffusion.set_index('ID')
    stats=net_analysis
    
    
    # calculate the final stats table that will be used for correlation and p-value measurements
    for i,v in int_diff_diffusion.iterrows():
        stats.at[i,'diffusion_centrality']=v[1]
        stats.at[i,'diffusion_vulnerability']=v[2]
        stats.at[i,'diffusion_noint']=v[3]
        stats.at[i,'diffusion_high_risk']=v[4]
#         stats.at[i,'diffusion_optimized']=v[5]
        
    for i,v in int_diff_contagion.iterrows():
        stats.at[i,'contagion_centrality']=v[1]
        stats.at[i,'contagion_vulnerability']=v[2]
        stats.at[i,'contagion_noint']=v[3]  
        stats.at[i,'contagion_high_risk']=v[4]
#         stats.at[i,'contagion_optimized']=v[5]
    
    print(stats)
    drop_list=[]
    if 'diffusion_centrality' in stats:
        drop_list.append('diffusion_centrality')
    if 'diffusion_vulnerability' in stats:
        drop_list.append('diffusion_vulnerability')        
    if 'diffusion_noint' in stats:
        drop_list.append('diffusion_noint') 
    if 'contagion_centrality' in stats:
        drop_list.append('contagion_centrality')  
    if 'contagion_vulnerability' in stats:
        drop_list.append('contagion_vulnerability')  
    if 'contagion_noint' in stats:
        drop_list.append('contagion_noint') 
    if 'diffusion_high_risk' in stats:
        drop_list.append('diffusion_high_risk')         
    if 'contagion_high_risk' in stats:
        drop_list.append('contagion_high_risk')  
#     if 'diffusion_optimized' in stats:
#         drop_list.append('diffusion_optimized')         
#     if 'contagion_optimized' in stats:
#         drop_list.append('contagion_optimized')         
        
    stats_corr = stats.corr()
    
    stats_corr=stats_corr.drop(drop_list)
    stats_corr=stats_corr[drop_list]
    stats_corr
    
    stats_pvalues = calculate_pvalues(stats)
    stats_pvalues = stats_pvalues.convert_objects(convert_numeric=True)
    stats_pvalues=stats_pvalues.drop(drop_list)
    stats_pvalues=stats_pvalues[drop_list]
    stats_pvalues
    
    # Generate a mask for the upper triangle
    mask = np.zeros_like(stats_corr, dtype=np.bool)
    mask[np.triu_indices_from(mask)] = True

    # Set up the matplotlib figure
    f, ax = plt.subplots(figsize=(11, 9))
    plt.title('Stats Correlations', fontsize=24)
    # Generate a custom diverging colormap
    cmap = sns.diverging_palette(220, 10, as_cmap=True)

    # Draw the heatmap with the mask and correct aspect ratio
    sns.heatmap(stats_corr, mask=mask, cmap=cmap, vmax=1.0, center=0,
                square=True, linewidths=.5, cbar_kws={"shrink": .5})    
    
    stats_pvalues_bin = stats_pvalues >= 0.01

    # Generate a mask for the upper triangle
    mask = np.zeros_like(stats_pvalues_bin, dtype=np.bool)
    mask[np.triu_indices_from(mask)] = True


    # Set up the matplotlib figure
    f, ax = plt.subplots(figsize=(11, 9))
    plt.title('Stats Significance (p<0.01)', fontsize=24)
    # Generate a custom diverging colormap
    cmap = sns.diverging_palette(220, 210, as_cmap=True)

    # Draw the heatmap with the mask and correct aspect ratio
    sns.heatmap(stats_pvalues_bin, mask=mask, cmap=cmap, vmax=1.0, center=0,
                square=True, linewidths=.5, cbar_kws={"shrink": .5})    
    
    

    if model=='all' or model=='diffusion':
        plt.figure(figsize=((20,20)))
#        plt.suptitle('Diffusion Model Scatter Plots')
#         plt.subplot(331)
#         x=stats[['RClosenessCentralization']].values
#         y=stats[['diffusion_centrality']].values
#         df = pd.DataFrame()
#         x=list(chain.from_iterable(x.tolist()))
#         y=list(chain.from_iterable(y.tolist()))
#         df['Closeness Centralization']=x
#         df['Most Connected']=y
#         ax = sns.regplot(x="Closeness Centralization", y="Most Connected", data=df)
# #         plt.scatter(x,y)
# #         plt.xlabel("Closeness Centralization")
# #         plt.ylabel("Centrality - Diffusion")


#         plt.subplot(332)
#         x=stats[['RClosenessCentralization']].values
#         y=stats[['diffusion_vulnerability']].values
#         df = pd.DataFrame()
#         x=list(chain.from_iterable(x.tolist()))
#         y=list(chain.from_iterable(y.tolist()))
#         df['Closeness Centralization']=x
#         df['Vulnerability']=y
#         ax = sns.regplot(x="Closeness Centralization", y="Vulnerability", data=df)        
# #         plt.scatter(x,y)
# #         plt.xlabel("Closeness Centralization")
# #         plt.ylabel("Vulnerability - Diffusion")

#         plt.subplot(333)
#         x=stats[['RClosenessCentralization']].values
#         y=stats[['diffusion_noint']].values
#         df = pd.DataFrame()
#         x=list(chain.from_iterable(x.tolist()))
#         y=list(chain.from_iterable(y.tolist()))
#         df['Closeness Centralization']=x
#         df['No Intervention']=y
#         ax = sns.regplot(x="Closeness Centralization", y="No Intervention", data=df) 
# #         plt.scatter(x,y)
# #         plt.xlabel("Closeness Centralization")
# #         plt.ylabel("No Intervention - Diffusion")


#         plt.subplot(334)
#         x=stats[['density']].values
#         y=stats[['diffusion_centrality']].values
#         df = pd.DataFrame()
#         x=list(chain.from_iterable(x.tolist()))
#         y=list(chain.from_iterable(y.tolist()))
#         df['Density']=x
#         df['Most Connected']=y
#         ax = sns.regplot(x="Density", y="Most Connected", data=df) 
# #         plt.scatter(x,y)
# #         plt.xlabel("Density")
# #         plt.ylabel("Centrality - Diffusion")


#         plt.subplot(335)
#         x=stats[['density']].values
#         y=stats[['diffusion_vulnerability']].values
#         df = pd.DataFrame()
#         x=list(chain.from_iterable(x.tolist()))
#         y=list(chain.from_iterable(y.tolist()))
#         df['Density']=x
#         df['Vulnerability']=y
#         ax = sns.regplot(x="Density", y="Vulnerability", data=df)        
# #         plt.scatter(x,y)
# #         plt.xlabel("Density")
# #         plt.ylabel("Vulnerability - Diffusion")

#         plt.subplot(336)
#         x=stats[['density']].values
#         y=stats[['diffusion_noint']].values
#         df = pd.DataFrame()
#         x=list(chain.from_iterable(x.tolist()))
#         y=list(chain.from_iterable(y.tolist()))
#         df['Density']=x
#         df['No Intervention']=y
#         ax = sns.regplot(x="Density", y="No Intervention", data=df) 
#         plt.scatter(x,y)
#         plt.xlabel("Density")
#         plt.ylabel("No Intervention - Diffusion")


        
        plt.subplot(337)
        x=stats[['ROutDegreeCentralization']].values
        y=stats[['diffusion_centrality']].values
        df = pd.DataFrame()
        x=list(chain.from_iterable(x.tolist()))
        y=list(chain.from_iterable(y.tolist()))
        df['Out Degree Centralization']=x
        df['Most Connected']=y
        ax = sns.regplot(x="Out Degree Centralization", y="Most Connected", data=df) 
#         plt.scatter(x,y)
#         plt.xlabel("OutDegree Centralization")
#         plt.ylabel("Centrality - Diffusion")


        plt.subplot(338)
        x=stats[['ROutDegreeCentralization']].values
        y=stats[['diffusion_vulnerability']].values
        df = pd.DataFrame()
        x=list(chain.from_iterable(x.tolist()))
        y=list(chain.from_iterable(y.tolist()))
        df['Out Degree Centralization']=x
        df['Vulnerability']=y
        ax = sns.regplot(x="Out Degree Centralization", y="Vulnerability", data=df) 
#         plt.scatter(x,y)
#         plt.xlabel("OutDegree Centralization")
#         plt.ylabel("Vulnerability - Diffusion")

        plt.subplot(339)
        x=stats[['ROutDegreeCentralization']].values
        y=stats[['diffusion_noint']].values
        df = pd.DataFrame()
        x=list(chain.from_iterable(x.tolist()))
        y=list(chain.from_iterable(y.tolist()))
        df['Out Degree Centralization']=x
        df['No Intervention']=y
        ax = sns.regplot(x="Out Degree Centralization", y="No Intervention", data=df) 
        3
        # more fancy plot
#         g = sns.jointplot("Out Degree Centralization", "No Intervention", data=df, kind="reg", size=8)

#         plt.scatter(x,y)
#         plt.xlabel("OutDegree Centralization")
#         plt.ylabel("No Intervention - Diffusion")
        
#         plt.subplot(340)
#         x=stats[['ROutDegreeCentralization']].values
#         y=stats[['diffusion_high_risk']].values
#         plt.scatter(x,y)
#         plt.xlabel("OutDegree Centralization")
#         plt.ylabel("High Risk - Diffusion")
        
#         plt.subplot(341)
#         x=stats[['RClosenessCentralization']].values
#         y=stats[['diffusion_high_risk']].values
#         plt.scatter(x,y)
#         plt.xlabel("RClosenessCentralization")
#         plt.ylabel("High Risk - Diffusion")

        
                
#         plt.subplot(342)
#         x=stats[['density']].values
#         y=stats[['diffusion_high_risk']].values
#         plt.scatter(x,y)
#         plt.xlabel("Density")
#         plt.ylabel("High Risk - Diffusion")
        
        plt.show()
    
    if model=='all' or model=='contagion':
        plt.figure(figsize=((18,18)))
        plt.suptitle('Contagion Model Scatter Plots')
        plt.subplot(331)
        x=stats[['RClosenessCentralization']].values
        y=stats[['contagion_centrality']].values
        plt.scatter(x,y)
        plt.xlabel("Closeness Centralization")
        plt.ylabel("Centrality - Contagion")


        plt.subplot(332)
        x=stats[['RClosenessCentralization']].values
        y=stats[['contagion_vulnerability']].values
        plt.scatter(x,y)
        plt.xlabel("Closeness Centralization")
        plt.ylabel("Vulnerability - Contagion")

        plt.subplot(333)
        x=stats[['RClosenessCentralization']].values
        y=stats[['contagion_noint']].values
        plt.scatter(x,y)
        plt.xlabel("Closeness Centralization")
        plt.ylabel("No Intervention - Contagion")


        plt.subplot(334)
        x=stats[['density']].values
        y=stats[['contagion_centrality']].values
        plt.scatter(x,y)
        plt.xlabel("Density")
        plt.ylabel("Centrality - Contagion")


        plt.subplot(335)
        x=stats[['density']].values
        y=stats[['contagion_vulnerability']].values
        plt.scatter(x,y)
        plt.xlabel("Density")
        plt.ylabel("Vulnerability - Contagion")

        plt.subplot(336)
        x=stats[['density']].values
        y=stats[['contagion_noint']].values
        plt.scatter(x,y)
        plt.xlabel("Density")
        plt.ylabel("No Intervention - Contagion")


        plt.subplot(337)
        x=stats[['ROutDegreeCentralization']].values
        y=stats[['contagion_centrality']].values
        plt.scatter(x,y)
        plt.xlabel("OutDegree Centralization")
        plt.ylabel("Centrality - Contagion")


        plt.subplot(338)
        x=stats[['ROutDegreeCentralization']].values
        y=stats[['contagion_vulnerability']].values
        plt.scatter(x,y)
        plt.xlabel("OutDegree Centralization")
        plt.ylabel("Vulnerability - Contagion")

        plt.subplot(339)
        x=stats[['ROutDegreeCentralization']].values
        y=stats[['contagion_noint']].values
        plt.scatter(x,y)
        plt.xlabel("OutDegree Centralization")
        plt.ylabel("No Intervention - Contagion")
        plt.show()    

    
    return stats_corr,stats_pvalues,stats

def fancy_heatmap(sr):
    # heatmap stuff
    # links :http://alanpryorjr.com/visualizations/seaborn/heatmap/heatmap/
    # https://www.quantinsti.com/blog/creating-heatmap-using-python-seaborn/ -> this is primary!
    success_rates=sr
    hm=success_rates[['ID','perc_sni']].sort_values('perc_sni',ascending=False)
    hm=hm.drop(hm.index[len(hm)-1])
    yrows=[1,1,1,1,1,2,2,2,2,2,3,3,3,3,3,4,4,4,4,4,5,5,5,5,5]
    hm=hm.assign(Yrows=yrows)
    xrows=[1,2,3,4,5,1,2,3,4,5,1,2,3,4,5,1,2,3,4,5,1,2,3,4,5]
    hm=hm.assign(Xrows=xrows)
    df_to_excel(hm, filename='../output/heatmap.xlsx', sheet='heatmap')
    hm=pd.ExcelFile('../output/heatmap.xlsx',sheet_name='heatmap')
    hm=hm.parse('heatmap')
    cls=(np.asarray(hm['ID'])).reshape(5,5)
    perc=(np.asarray(hm['perc_sni'])).reshape(5,5)
    result=hm.pivot(index='Yrows',columns='Xrows',values='perc_sni')
    labels=(np.asarray(["Class {0} \n \n {1:.2f}%".format(c,p)
                      for c,p in zip(cls.flatten(),perc.flatten())])
           ).reshape(5,5)
    fig, ax =plt.subplots(figsize=(15,10))
    title= "Success Rates per Class"
    plt.title(title,fontsize=18)
    ttl=ax.title
    ttl.set_position([0.5,1.05])
    ax.set_xticks([])
    ax.set_yticks([])
    ax.axis('off')
    sns.heatmap(result,annot=labels,annot_kws={"size": 16},fmt="",cmap="Blues",linewidths=0.30,ax=ax)
    plt.show()
    # def WriteDictToCSV(csv_file,csv_columns,dict_data):
#     try:
#         with open(csv_file, 'w') as csvfile:
#             writer = csv.DictWriter(csvfile, fieldnames=csv_columns)
#             writer.writeheader()
#             for data in dict_data:
#                 writer.writerow(data)
#     except IOError:
#         print("I/O error", csv_file)
#     return          